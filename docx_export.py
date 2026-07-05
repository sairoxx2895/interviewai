"""
docx_export.py
---------------
Turns a saved session (role, questions, flags, scoring guides...) into a
clean, print-ready Microsoft Word document using python-docx.
"""

import os
import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

EXPORT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "exports")
os.makedirs(EXPORT_DIR, exist_ok=True)

BRAND = RGBColor(0x1C, 0x2B, 0x3A)      # ink navy
GREEN = RGBColor(0x2E, 0x7D, 0x4F)      # green flag
RED = RGBColor(0xB5, 0x47, 0x3F)        # red flag
MUTED = RGBColor(0x5B, 0x66, 0x70)      # muted gray


def _slugify(text: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", text.strip().lower()).strip("-")
    return slug or "session"


def _add_bullets(doc, items, color):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.color.rgb = color
        run.font.size = Pt(10.5)


def build_docx(session: dict) -> str:
    doc = Document()

    # Base typography
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- Title block -----------------------------------------------------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("InterviewAI")
    run.font.size = Pt(14)
    run.font.color.rgb = MUTED
    run.bold = True

    heading = doc.add_paragraph()
    run = heading.add_run(f"{session.get('role', 'Interview')} — Interview Question Set")
    run.font.size = Pt(22)
    run.font.color.rgb = BRAND
    run.bold = True

    meta = doc.add_paragraph()
    meta_text = (
        f"Experience level: {session.get('experience_level', '—')}    |    "
        f"Interview type: {session.get('interview_type', '—')}    |    "
        f"Questions: {len(session.get('questions', []))}"
    )
    run = meta.add_run(meta_text)
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED
    run.italic = True

    created = session.get("created_at")
    if created:
        date_p = doc.add_paragraph()
        run = date_p.add_run(f"Generated: {created}")
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED

    if session.get("overall_evaluation_notes"):
        doc.add_paragraph()
        note_heading = doc.add_paragraph()
        run = note_heading.add_run("General Guidance for the Interviewer")
        run.bold = True
        run.font.color.rgb = BRAND
        doc.add_paragraph(session["overall_evaluation_notes"])

    doc.add_paragraph("─" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Questions ---------------------------------------------------------
    for idx, q in enumerate(session.get("questions", []), start=1):
        qh = doc.add_paragraph()
        run = qh.add_run(f"Q{idx}. {q.get('question', '')}")
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = BRAND

        cat = doc.add_paragraph()
        run = cat.add_run(f"Category: {q.get('category', '—')}")
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED
        run.italic = True

        if q.get("follow_up"):
            fu = doc.add_paragraph()
            run = fu.add_run("Follow-up: ")
            run.bold = True
            run.font.size = Pt(10.5)
            run2 = fu.add_run(q["follow_up"])
            run2.font.size = Pt(10.5)

        if q.get("evaluation_tip"):
            et = doc.add_paragraph()
            run = et.add_run("What a strong answer looks like: ")
            run.bold = True
            run.font.size = Pt(10.5)
            run2 = et.add_run(q["evaluation_tip"])
            run2.font.size = Pt(10.5)

        if q.get("green_flags"):
            gf = doc.add_paragraph()
            run = gf.add_run("Green flags")
            run.bold = True
            run.font.color.rgb = GREEN
            run.font.size = Pt(10.5)
            _add_bullets(doc, q["green_flags"], GREEN)

        if q.get("red_flags"):
            rf = doc.add_paragraph()
            run = rf.add_run("Red flags")
            run.bold = True
            run.font.color.rgb = RED
            run.font.size = Pt(10.5)
            _add_bullets(doc, q["red_flags"], RED)

        scoring = q.get("scoring_guide") or {}
        if scoring:
            sg = doc.add_paragraph()
            run = sg.add_run("Scoring guide (1-5)")
            run.bold = True
            run.font.size = Pt(10.5)
            run.font.color.rgb = BRAND
            labels = {"1_2": "1-2 (Weak)", "3": "3 (Adequate)", "4_5": "4-5 (Strong)"}
            for key, label in labels.items():
                if scoring.get(key):
                    p = doc.add_paragraph(style="List Bullet 2")
                    run = p.add_run(f"{label}: {scoring[key]}")
                    run.font.size = Pt(10)

        if idx != len(session.get("questions", [])):
            doc.add_paragraph("· · ·").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Footer --------------------------------------------------------
    doc.add_paragraph()
    footer = doc.add_paragraph()
    run = footer.add_run("Generated with InterviewAI — structured, objective, bias-aware interviewing.")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    run.italic = True

    filename = f"{_slugify(session.get('role', 'interview'))}-{datetime.utcnow().strftime('%Y%m%d-%H%M%S')}.docx"
    path = os.path.join(EXPORT_DIR, filename)
    doc.save(path)
    return path
